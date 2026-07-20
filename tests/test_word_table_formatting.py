from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm

from src.word_report import _table


def test_table_cells_override_normal_indent_and_apply_bilingual_centered_formatting():
    document = Document()
    document.styles["Normal"].paragraph_format.first_line_indent = Cm(0.74)
    headers = ["变量", "原始缺测", "零值", "硬范围", "Hampel", "恒定值", "人工删除", "最终有效"]
    _table(document, "表1-2 站点质控汇总", headers, [("  多环芳烃\t", "1", "2", "3", "4", "5", "6", "7")])

    table = document.tables[0]
    assert [cell.text for cell in table.rows[0].cells] == headers
    assert table.cell(1, 0).text == "多环芳烃"
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            assert cell.vertical_alignment == WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                assert paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
                assert paragraph.paragraph_format.first_line_indent == 0
                assert paragraph.paragraph_format.left_indent == 0
                assert paragraph.paragraph_format.right_indent == 0
                assert paragraph.paragraph_format.space_before == 0
                assert paragraph.paragraph_format.space_after == 0
                assert paragraph.paragraph_format.line_spacing == 1
                assert not paragraph.text.startswith((" ", "\t", "　"))
                assert "\t" not in paragraph.text
                for run in paragraph.runs:
                    assert run.bold is (row_index == 0)
                    assert run.font.size.pt == 9
                    fonts = run._element.rPr.rFonts
                    assert fonts.get(qn("w:eastAsia")) == "宋体"
                    assert fonts.get(qn("w:ascii")) == "Times New Roman"
                    assert fonts.get(qn("w:hAnsi")) == "Times New Roman"
                    assert fonts.get(qn("w:cs")) == "Times New Roman"

    # Body paragraphs still retain the station report's two-character indent.
    assert document.styles["Normal"].paragraph_format.first_line_indent > 0
