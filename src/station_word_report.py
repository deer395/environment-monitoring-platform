"""Deterministic station-level Word report rendering."""

from io import BytesIO

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from .word_report import (
    _add_heading, _add_picture, _anomaly_figure, _comparison_figure, _date,
    _fmt, _hourly_daily_figure, _monthly_figure, _page_number, _qc_text,
    _table, _value_unit,
)
from .manual_qc import summarize_candidate_decisions


def _set_run_fonts(run, east_asia="宋体", latin="Times New Roman", bold=None, size=None):
    run.font.name = latin
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for name, value in (("eastAsia", east_asia), ("ascii", latin), ("hAnsi", latin), ("cs", latin)):
        fonts.set(qn(f"w:{name}"), value)
    if bold is not None: run.bold = bold
    if size is not None: run.font.size = Pt(size)


def _configure_styles(doc):
    normal = doc.styles["Normal"]
    # Style rPr needs the same four explicit OOXML mappings as ordinary runs.
    fonts = normal.element.get_or_add_rPr().get_or_add_rFonts()
    for name, value in (("eastAsia", "宋体"), ("ascii", "Times New Roman"), ("hAnsi", "Times New Roman"), ("cs", "Times New Roman")): fonts.set(qn(f"w:{name}"), value)
    normal.font.size = Pt(12); normal.font.bold = False
    normal.paragraph_format.first_line_indent = Cm(.74); normal.paragraph_format.line_spacing = 1.5; normal.paragraph_format.space_after = Pt(4)
    for style_name, size in (("Heading 1", 15), ("Heading 2", 14)):
        style = doc.styles[style_name]; fonts = style.element.get_or_add_rPr().get_or_add_rFonts()
        for name, value in (("eastAsia", "黑体"), ("ascii", "Times New Roman"), ("hAnsi", "Times New Roman"), ("cs", "Times New Roman")): fonts.set(qn(f"w:{name}"), value)
        style.font.size = Pt(size); style.font.bold = True; style.paragraph_format.first_line_indent = Cm(0)


def _heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level); paragraph.paragraph_format.first_line_indent = Cm(0); return paragraph


def _body(doc, text):
    paragraph = doc.add_paragraph(text); paragraph.paragraph_format.first_line_indent = Cm(.74)
    for run in paragraph.runs: _set_run_fonts(run, bold=False, size=12)
    return paragraph


def _date_only(value):
    value = pd.Timestamp(value)
    return f"{value.year}年{value.month}月{value.day}日"


def _overview_figure(variables, title):
    fig, axes = plt.subplots(len(variables), 1, figsize=(7.1, 8.4), sharex=True)
    for axis, item in zip(axes, variables):
        context = item["context"]
        name, unit = context["variable_info"]["display_name"], context["variable_info"]["unit"]
        daily = context["time_features"]["resampled"]["daily"]
        axis.plot(daily.get("datetime"), daily.get("value"), linewidth=.8, color="#1f77b4")
        axis.set_ylabel(f"{name}\n({unit})", fontsize=10)
        axis.tick_params(labelsize=9)
        axis.grid(True, alpha=.3)
    axes[-1].set_xlabel("时间", fontsize=10); fig.suptitle(title, fontsize=11)
    fig.autofmt_xdate(); output = BytesIO(); fig.tight_layout(rect=(0, 0, 1, .98)); fig.savefig(output, format="png", dpi=300, bbox_inches="tight"); plt.close(fig); output.seek(0)
    return output


def _station_qc_text(data, qc):
    auto_zero, auto_range = qc.get("removed_by_sensor_zero", 0), qc.get("removed_by_range", 0)
    auto = auto_zero + auto_range
    parts = []
    if auto: parts.append(f"自动规则删除{auto:,}条记录，其中包括{auto_zero:,}条传感器零值和{auto_range:,}条超出硬范围的记录。")
    candidates = []
    if qc.get("flagged_by_hampel", 0): candidates.append(f"Hampel规则识别{qc['flagged_by_hampel']:,}条候选记录")
    if qc.get("flagged_by_constant_value", 0): candidates.append(f"连续恒定值规则识别{qc['flagged_by_constant_value']:,}条候选记录")
    if candidates: parts.append("，".join(candidates) + "。")
    review = qc.get("_review_table")
    if review is not None and not review.empty:
        summary = summarize_candidate_decisions(review); removed, kept, total = summary["candidate_removed_count"], summary["candidate_kept_count"], summary["unique_candidate_count"]
        if total and removed == total: parts.append(f"经人工复核，上述{total:,}条候选记录均予以删除。")
        elif total and kept == total: parts.append(f"经人工复核，上述{total:,}条候选记录均予以保留。")
        elif total: parts.append(f"经人工复核，其中{removed:,}条予以删除，{kept:,}条予以保留。")
    else:
        manual = qc.get("manual_remove_count", 0)
        if manual: parts.append(f"经人工复核，另有{manual:,}条记录被删除。")
    parts.append(f"最终获得{data['final_valid_count']:,}条有效记录，数据有效率为{data['valid_rate']:.2f}%。")
    return "".join(parts)


def _variable_section(doc, item, index, figure_number):
    context = item["context"]
    variable, data, qc, stats, features = (context[key] for key in ("variable_info", "data_summary", "qc_summary", "statistics", "time_features"))
    name, unit = variable["display_name"], variable["unit"]
    _heading(doc, f"3.{index} {name}", 2)
    qc = {**qc, "_review_table": context["source_data"].get("review_table")}
    _body(doc, f"{name}数据分析时段为{_date_only(data['start_time'])}至{_date_only(data['end_time'])}，共包含{data['raw_count']:,}条原始记录。" + _station_qc_text(data, qc))
    source = context["source_data"]
    _add_picture(doc, _comparison_figure(source["raw_data"], source["final_qc_data"], name, unit), f"图3-{figure_number} {name}原始数据与最终质控数据对比")
    interp = features["interpretation"]
    _table(doc, f"表3-{index} {name}基础统计", ["指标", "数值", "单位"], [("最小值", _fmt(stats.get("min")), unit), ("最大值", _fmt(stats.get("max")), unit), ("平均值", _fmt(stats.get("mean")), unit), ("中位数", _fmt(stats.get("median")), unit), ("标准差", _fmt(stats.get("std")), unit), ("P5", _fmt(interp.get("p5")), unit), ("P95", _fmt(interp.get("p95")), unit)])
    for key in ("hourly", "intraday", "monthly"):
        text = context["narratives"].get(key, "")
        text = text.replace("日平均序列整体存在阶段性起伏。", "").replace("日平均序列呈持续性变化。", "").replace("日平均序列在前期总体维持较高水平，后期出现明显水平变化。", "")
        _body(doc, text)
    resampled, anomaly = features["resampled"], features["anomaly"]
    _add_picture(doc, _hourly_daily_figure(resampled["hourly"], resampled["daily"], name, unit), f"图3-{figure_number + 1} {name}小时平均与日平均")
    _add_picture(doc, _anomaly_figure(anomaly, name, unit), f"图3-{figure_number + 2} {name}日内距平")
    _add_picture(doc, _monthly_figure(stats.get("monthly"), name, unit), f"图3-{figure_number + 3} {name}月平均与月标准差")


def generate_station_word_report(context):
    """Render a station report from the already validated station context."""
    doc = Document(); section = doc.sections[0]; section.left_margin = section.right_margin = Cm(2.2)
    _configure_styles(doc)
    footer = section.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.CENTER; _page_number(footer)
    info, variables = context["station_info"], context["variables"]
    doc.add_paragraph(info.get("report_title") or "站点环境监测综合报告", style="Title").alignment = WD_ALIGN_PARAGRAPH.CENTER
    for label, value in (("项目名称", info.get("project_name")), ("站点名称", info.get("site_name")), ("编制部门", info.get("department")), ("编制人", info.get("author"))):
        if value: _body(doc, f"{label}：{value}").paragraph_format.first_line_indent = Cm(0)
    doc.add_paragraph("报告生成日期：" + _date_only(variables[0]["context"]["software_info"]["generated_at"])); doc.add_page_break()
    _heading(doc, "1. 站点数据与质控汇总")
    _table(doc, "表1-1 站点数据概况", ["变量", "单位", "分析时间范围", "原始记录", "最终有效", "最终缺测", "有效率"], context["overview_rows"])
    _table(doc, "表1-2 站点质控汇总", ["变量", "原始缺测", "零值", "硬范围", "Hampel", "恒定值", "人工删除", "最终有效"], context["qc_rows"])
    _body(doc, "Hampel和恒定值列表示算法识别的候选记录数，候选记录是否删除以人工复核结果为准。")
    _body(doc, "本报告中的统计结果和图件均基于人工确认后的最终质控数据，主要用于描述各监测要素的数据质量、统计特征和时间变化，不涉及环境质量等级或变化成因评价。")
    _heading(doc, "2. 多要素时间序列总览")
    _add_picture(doc, _overview_figure(variables[:4], "水文与基础环境要素日平均时间序列"), "图2-1 水文与基础环境要素日平均时间序列")
    _add_picture(doc, _overview_figure(variables[4:], "水质与生态要素日平均时间序列"), "图2-2 水质与生态要素日平均时间序列")
    _body(doc, "各曲线均基于最终质控数据计算日平均值，无有效日平均数据的时段保持为空，未进行插值处理。")
    _heading(doc, "3. 各监测要素质控与统计分析")
    for index, item in enumerate(variables, 1): _variable_section(doc, item, index, (index - 1) * 4 + 1)
    output = BytesIO(); doc.save(output); return output.getvalue()
