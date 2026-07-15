"""Lightweight V4.1 report validation; uses synthetic data only."""

from io import BytesIO
from pathlib import Path
import sys

import pandas as pd
from docx import Document

PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from src.anomaly import calculate_configured_anomaly
from src.manual_qc import apply_review_table_decisions, build_qc_review_table
from src.metrics import calculate_metrics
from src.qc import apply_quality_control
from src.report_context import build_report_context
from src.resampling import resample_configured
from src.variable_registry import get_variable_metadata
from src.word_report import generate_single_variable_report


def check(condition, message):
    if not condition:
        raise AssertionError(message)


def make_context(variable_key="depth", token="confirmed-token"):
    index = pd.date_range("2026-01-01", periods=96, freq="h")
    raw = pd.DataFrame({"record_id": [f"r{i}" for i in range(len(index))], "datetime": index, "value": [5 + (i % 12) * .1 for i in range(len(index))]})
    raw.loc[4, "value"] = 0.0
    raw.loc[5, "value"] = 999.0
    raw.attrs.update({"variable_key": variable_key, "unit": get_variable_metadata(variable_key)["unit"]})
    metadata = get_variable_metadata(variable_key)
    auto, summary, qc_log = apply_quality_control(raw, metadata, enable_hampel=True, enable_constant_value=True)
    review = build_qc_review_table(raw, auto, qc_log)
    review.loc[10, "user_decision"] = "manual_remove"
    final, final_log = apply_review_table_decisions(raw, auto, qc_log, review)
    resampled = resample_configured(final, metadata)
    anomaly = calculate_configured_anomaly(resampled, metadata)
    metrics = calculate_metrics(variable_key, resampled["hourly"], resampled["daily"], anomaly, metadata=metadata, base_data=resampled["hourly"])
    return build_report_context(variable_key=variable_key, raw_data=raw, final_qc_data=final, final_qc_log=final_log, qc_summary=summary, review_table=review, resampled=resampled, anomaly=anomaly, metrics=metrics, qc_token=token, confirmed_qc_token=token, project_name="测试项目", author="测试人员")


def main():
    context = make_context()
    report = generate_single_variable_report(context)
    check(len(report) > 10_000, "Word report should not be empty")
    document = Document(BytesIO(report))
    text = "\n".join(p.text for p in document.paragraphs)
    for required in ("水深", "m", "数据摘要", "数据质控与前后对比", "统计特征", "结论", "附录A", "附录C"):
        check(required in text, f"report missing {required}")
    table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    check("P5" in table_text and "日变化幅度P95" in table_text, "statistics table missing quantiles")
    check("重复时间" not in text, "report must not show duplicate-time content")
    for internal_label in ("increase", "decrease", "stable", "change_point", "high_variability", "periodic", "2026-05", "NaN", "主要变化特征为："):
        check(internal_label not in text, f"report leaked internal text: {internal_label}")
    check("传感器零值" not in text.split("2. 数据质控与前后对比")[0], "data summary must not contain QC counts")
    analysis_index = text.index("时间变化综合分析")
    hourly_index = text.index("小时平均", analysis_index)
    intraday_index = text.index("日内距平", analysis_index)
    monthly_index = text.index("月平均", intraday_index)
    figure2_index = text.index("图2", monthly_index)
    figure3_index = text.index("图3", figure2_index)
    figure4_index = text.index("图4", figure3_index)
    check(analysis_index < hourly_index < intraday_index < monthly_index < figure2_index < figure3_index < figure4_index, "narratives and figures must follow hourly-intraday-monthly order")
    check(len(document.tables) >= 5, "report should contain required tables")
    check(len(document.inline_shapes) >= 4, "report should contain report figures")
    try:
        broken = dict(context); broken["software_info"] = dict(context["software_info"], qc_confirmed=False)
        generate_single_variable_report(broken)
        raise AssertionError("unconfirmed report generation must fail")
    except ValueError:
        pass
    try:
        build_report_context(variable_key="depth", raw_data=pd.DataFrame(), final_qc_data=pd.DataFrame(), final_qc_log=None, qc_summary={}, review_table=None, resampled={}, anomaly=pd.DataFrame(), metrics={}, qc_token="new", confirmed_qc_token="old")
        raise AssertionError("changed qc token must fail")
    except ValueError:
        pass
    pahs_context = make_context("pahs")
    check(pahs_context["variable_info"]["unit"] == "ppb", "PAHs unit must be ppb")
    pahs_context["statistics"]["monthly"] = pd.DataFrame(columns=["year_month", "monthly_mean", "monthly_std"])
    check(len(generate_single_variable_report(pahs_context)) > 10_000, "empty monthly data must not fail")
    print("V4.1 single-variable Word report preview passed")


if __name__ == "__main__":
    main()
