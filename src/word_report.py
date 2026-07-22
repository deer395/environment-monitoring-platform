"""Deterministic Word report and report-only matplotlib figures."""

from __future__ import annotations

from io import BytesIO
import math

import matplotlib
matplotlib.use("Agg")
import matplotlib.dates as mdates
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from .time_series_interpretation import MODE_LABELS


plt.rcParams.update({"font.sans-serif": ["Microsoft YaHei", "SimHei", "Arial Unicode MS", "DejaVu Sans"], "axes.unicode_minus": False})


def _fmt(value, digits=None):
    if value is None or pd.isna(value):
        return "—"
    if isinstance(value, (int, float)) and not isinstance(value, bool):
        magnitude = abs(value)
        digits = digits if digits is not None else (2 if magnitude >= 1 else 3 if magnitude >= .01 else 4)
        return f"{value:.{digits}f}".rstrip("0").rstrip(".")
    return str(value)


def _count(value):
    return f"{int(value):,}"


def _value_unit(value, unit):
    return f"{_fmt(value)} {unit}" if pd.notna(value) else "—"


def _range_unit(low, high, unit):
    return f"{_fmt(low)}～{_fmt(high)} {unit}" if pd.notna(low) and pd.notna(high) else "—"


def _month(value):
    period = pd.Period(str(value), freq="M")
    return f"{period.year}年{period.month}月"


def _month_span(start, end):
    return _month(start) if start == end else f"{_month(start)}至{_month(end)}"


def _date(value):
    value = pd.to_datetime(value, errors="coerce")
    if pd.isna(value):
        return "—"
    return f"{value.year}年{value.month}月{value.day}日{value.hour}时" if value.hour or value.minute else f"{value.year}年{value.month}月{value.day}日"


def _pct(count, total):
    return "无法计算" if not total else f"{count / total * 100:.2f}%"


def _qc_text(data, qc):
    sentences = [f"原始数据共{_count(data['raw_count'])}条。"]
    auto = []
    if qc.get("removed_by_sensor_zero", 0): auto.append(f"{_count(qc['removed_by_sensor_zero'])}条传感器零值")
    if qc.get("removed_by_range", 0): auto.append(f"{_count(qc['removed_by_range'])}条超出硬范围的记录")
    if auto: sentences.append("自动规则识别并剔除" + "和".join(auto) + "。")
    candidates = []
    if qc.get("flagged_by_hampel", 0): candidates.append(f"Hampel规则识别{_count(qc['flagged_by_hampel'])}条候选记录")
    if qc.get("flagged_by_constant_value", 0): candidates.append(f"连续恒定值规则识别{_count(qc['flagged_by_constant_value'])}条候选记录")
    if candidates: sentences.append("，".join(candidates) + "，以上候选均进入人工复核环节。")
    ending = f"最终形成{_count(data['final_valid_count'])}条有效记录，有效率为{_fmt(data['valid_rate'], 2)}%。"
    sentences.append((f"经人工复核后，共有{_count(qc['manual_remove_count'])}条记录被删除，" if qc.get("manual_remove_count", 0) else "") + ending)
    return "".join(sentences)


def _comparison_text(source, interpretation, unit):
    raw = pd.to_numeric(source["raw_data"].get("value"), errors="coerce")
    final = pd.to_numeric(source["final_qc_data"].get("value"), errors="coerce")
    p5, p95 = interpretation.get("p5"), interpretation.get("p95")
    high = raw.notna().any() and pd.notna(p95) and raw.max() > p95 + max((p95 - p5) * .5, 1e-9)
    low = raw.notna().any() and pd.notna(p5) and raw.min() < p5 - max((p95 - p5) * .5, 1e-9)
    opening = "原始时间序列中存在多次短时高值尖峰，同时伴有少量异常低值。" if high and low else "原始时间序列中存在多次短时高值尖峰。" if high else "原始时间序列中存在少量异常低值。" if low else "质控前后数据主体分布变化有限。"
    if final.notna().any() and pd.notna(p5) and pd.notna(p95):
        return opening + f"经自动规则和人工复核后，上述异常记录被处理。最终质控后的原始记录范围为{_range_unit(final.min(), final.max(), unit)}；经小时平均后，90%的数据分布在{_range_unit(p5, p95, unit)}之间，主要时间变化过程得到保留。"
    return opening


REPORT_FIGURE_DPI = 300
REPORT_FIGURE_FACE = "#ffffff"
REPORT_GRID_COLOR = "#d9e0e6"
REPORT_RAW_COLOR = "#425a70"
REPORT_FINAL_COLOR = "#a33a3a"


def _style_report_axis(ax, title, ylabel, xlabel="时间", legend=False):
    """Apply print-oriented visual styling without changing plotted data."""
    ax.set_facecolor(REPORT_FIGURE_FACE)
    ax.set_title(title, fontsize=15, fontweight="semibold", color="#173b5c", pad=10, loc="left")
    ax.set_ylabel(ylabel, fontsize=11, color="#273746")
    if xlabel:
        ax.set_xlabel(xlabel, fontsize=11, color="#273746")
    ax.tick_params(axis="both", labelsize=9, colors="#4d5b67")
    ax.grid(True, color=REPORT_GRID_COLOR, linewidth=.65, alpha=.9)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_color("#aebbc6")
    ax.spines["bottom"].set_color("#aebbc6")
    if legend:
        ax.legend(frameon=False, fontsize=9, loc="upper right")


def _figure(fig):
    output = BytesIO()
    fig.patch.set_facecolor(REPORT_FIGURE_FACE)
    fig.tight_layout(pad=1.2)
    fig.savefig(output, format="png", dpi=REPORT_FIGURE_DPI, bbox_inches="tight", facecolor=REPORT_FIGURE_FACE)
    plt.close(fig)
    output.seek(0)
    return output


def _comparison_figure(raw, final, name, unit):
    """Keep the original/final two-panel evidence view with aligned time bounds."""
    fig, axes = plt.subplots(2, 1, figsize=(10.5, 7.2), sharex=True)
    raw_values = pd.to_numeric(raw.get("value"), errors="coerce")
    final_values = pd.to_numeric(final.get("value"), errors="coerce")
    raw_times = pd.to_datetime(raw.get("datetime"), errors="coerce")
    final_times = pd.to_datetime(final.get("datetime"), errors="coerce")
    axes[0].plot(raw_times, raw_values, color=REPORT_RAW_COLOR, linewidth=.75, label="原始数据")
    axes[1].plot(final_times, final_values, color=REPORT_FINAL_COLOR, linewidth=.85, label="最终质控数据")
    _style_report_axis(axes[0], "原始数据", f"{name} ({unit})", xlabel="", legend=True)
    _style_report_axis(axes[1], "最终质量控制数据", f"{name} ({unit})", legend=True)
    all_times = pd.concat([raw_times, final_times]).dropna()
    if not all_times.empty:
        axes[0].set_xlim(all_times.min(), all_times.max())
    axes[1].xaxis.set_major_locator(mdates.AutoDateLocator())
    axes[1].xaxis.set_major_formatter(mdates.DateFormatter("%Y/%m/%d"))
    fig.autofmt_xdate(rotation=20, ha="right")
    fig.subplots_adjust(hspace=.34)
    return _figure(fig)


def _hourly_daily_figure(hourly, daily, name, unit):
    fig, ax = plt.subplots(figsize=(10.5, 4.6))
    ax.plot(hourly.get("datetime"), hourly.get("value"), label="小时平均", linewidth=.8, color=REPORT_RAW_COLOR)
    ax.plot(daily.get("datetime"), daily.get("value"), label="日平均", linewidth=1.5, color=REPORT_FINAL_COLOR)
    _style_report_axis(ax, "小时平均与日平均", f"{name} ({unit})", legend=True)
    fig.autofmt_xdate(rotation=20, ha="right")
    return _figure(fig)


def _anomaly_figure(anomaly, name, unit):
    fig, ax = plt.subplots(figsize=(10.5, 4.6))
    ax.plot(anomaly.get("datetime"), anomaly.get("anomaly"), linewidth=.8, color=REPORT_RAW_COLOR)
    ax.axhline(0, color="#71808d", linewidth=.8)
    _style_report_axis(ax, "日内距平", f"{name}日内距平 ({unit})")
    fig.autofmt_xdate(rotation=20, ha="right")
    return _figure(fig)


def _monthly_figure(monthly, name, unit):
    fig, ax = plt.subplots(figsize=(10.5, 4.6))
    data = monthly.dropna(subset=["monthly_mean"]) if monthly is not None else pd.DataFrame()
    if not data.empty:
        ax.errorbar(data["year_month"].astype(str), data["monthly_mean"], yerr=data["monthly_std"].fillna(0), marker="o", markersize=4, capsize=3, linewidth=1.1, color=REPORT_FINAL_COLOR)
    _style_report_axis(ax, "月平均与标准差", f"{name} ({unit})", xlabel="年月")
    ax.tick_params(axis="x", rotation=20)
    return _figure(fig)

def _format_body_paragraph(paragraph, font_size, *, line_spacing=None, space_after=None):
    """Format ordinary report prose with a two-Chinese-character first-line indent."""
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Pt(font_size * 2)
    fmt.left_indent = Pt(0)
    fmt.right_indent = Pt(0)
    if line_spacing is not None:
        fmt.line_spacing = line_spacing
    if space_after is not None:
        fmt.space_after = space_after
    return paragraph

def _set_table_run_fonts(run, bold=False, size=9):
    """Apply the report's bilingual font mapping to a table run."""
    run.font.name = "Times New Roman"
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for name, value in (("eastAsia", "宋体"), ("ascii", "Times New Roman"), ("hAnsi", "Times New Roman"), ("cs", "Times New Roman")):
        fonts.set(qn(f"w:{name}"), value)
    run.bold = bold
    run.font.size = Pt(size)


def format_table_cell(
    cell,
    *,
    bold=False,
    font_size=9,
    horizontal_center=True,
    vertical_center=True,
):
    """Make a table cell independent from the Normal paragraph style.

    Station reports intentionally give Normal paragraphs a first-line indent.
    Every table paragraph must explicitly override that setting; centering alone
    does not prevent the inherited indent from shifting or clipping its text.
    """
    if vertical_center:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        # Tables never use tabs for layout.  Remove any accidental layout
        # whitespace while retaining ordinary spaces within a value or date.
        cleaned = paragraph.text.replace("\t", "").strip(" \t\u3000")
        if cleaned != paragraph.text:
            paragraph.clear()
            paragraph.add_run(cleaned)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if horizontal_center else WD_ALIGN_PARAGRAPH.LEFT
        paragraph_format = paragraph.paragraph_format
        paragraph_format.first_line_indent = Pt(0)
        paragraph_format.left_indent = Pt(0)
        paragraph_format.right_indent = Pt(0)
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)
        paragraph_format.line_spacing = 1
        paragraph_format.tab_stops.clear_all()
        for run in paragraph.runs:
            _set_table_run_fonts(run, bold=bold, size=font_size)


def _set_cell_text(cell, text, bold=False):
    # Strip only leading/trailing layout whitespace.  Internal spaces, such as
    # those in dates and value/unit pairs, remain part of the report content.
    cell.text = str(text).replace("\t", "").strip(" \t\u3000")
    format_table_cell(cell, bold=bold)


def _table(document, title, headers, rows):
    document.add_paragraph(title).alignment = WD_ALIGN_PARAGRAPH.CENTER
    table = document.add_table(rows=1, cols=len(headers)); table.style = "Table Grid"
    for cell, value in zip(table.rows[0].cells, headers): _set_cell_text(cell, value, True)
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row): _set_cell_text(cell, value)
    document.add_paragraph()


def _caption(document, text):
    paragraph = document.add_paragraph(text); paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _page_number(paragraph):
    run = paragraph.add_run(); fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE"); run._r.addnext(fld)


def _add_picture(document, image, caption):
    document.add_picture(image, width=Cm(16)); document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER; _caption(document, caption)


def _add_heading(document, text):
    document.add_heading(text, level=1)


def _analysis_paragraphs(interp, stats, unit):
    monthly = interp.get("monthly_interpretation", {})
    paragraphs = []
    if monthly.get("segments"):
        segments = monthly["segments"]
        first = segments[0]
        text = f"{_month_span(first['start_month'], first['end_month'])}，月平均值总体"
        text += "维持在相对稳定水平。" if first.get("direction") == "stable" else ("呈下降变化。" if first.get("direction") == "decrease" else "呈上升变化。")
        if monthly.get("change_month"):
            direction = MODE_LABELS.get(monthly.get("direction"), "变化")
            change = f"{_month(monthly['change_month'])}起出现明显{direction}。"
            if pd.notna(monthly.get("pre_mean")) and pd.notna(monthly.get("post_mean")):
                change += f"转折前月平均水平约为{_value_unit(monthly['pre_mean'], unit)}，转折后约为{_value_unit(monthly['post_mean'], unit)}，{direction}约{_value_unit(monthly['absolute_change'], unit)}"
                if pd.notna(monthly.get("relative_change_pct")): change += f"，幅度约{_fmt(monthly['relative_change_pct'], 1)}%"
                change += "。"
            text += change
        paragraphs.append(text)
    if pd.notna(interp.get("max_value")) and pd.notna(interp.get("min_value")):
        paragraphs.append(f"最高小时平均值出现在{_date(interp['max_time'])}，为{_value_unit(interp['max_value'], unit)}；最低小时平均值出现在{_date(interp['min_time'])}，为{_value_unit(interp['min_value'], unit)}。90%的小时平均值分布在{_range_unit(interp.get('p5'), interp.get('p95'), unit)}之间。")
    month_sentence = []
    if monthly.get("monthly_max_month") is not None: month_sentence.append(f"月平均最高值出现在{_month(monthly['monthly_max_month'])}，为{_value_unit(monthly['monthly_max'], unit)}")
    if monthly.get("monthly_min_month") is not None: month_sentence.append(f"最低值出现在{_month(monthly['monthly_min_month'])}，为{_value_unit(monthly['monthly_min'], unit)}")
    if monthly.get("max_std_month") is not None: month_sentence.append(f"{_month(monthly['max_std_month'])}月标准差最大，为{_value_unit(monthly['max_std'], unit)}，该月数据波动相对更明显")
    if month_sentence: paragraphs.append("；".join(month_sentence) + "。")
    if pd.notna(interp.get("daily_range_p95")):
        periodic = interp.get("primary_mode") == "periodic"
        if periodic:
            detail = interp.get("periodicity", {})
            text = "日内距平呈持续而规律的重复波动"
            if pd.notna(detail.get("period_hours")): text += f"，近似周期约为{_fmt(detail['period_hours'], 1)}小时"
            if pd.notna(detail.get("autocorrelation")): text += f"，固定滞后相关系数为{_fmt(detail['autocorrelation'], 2)}"
            text += "。"
        else:
            text = f"日内距平整体围绕0 {unit}波动。"
        text += f"日变化幅度中位数为{_value_unit(interp.get('daily_range_median'), unit)}，95%的日变化幅度不超过{_value_unit(interp.get('daily_range_p95'), unit)}"
        if pd.notna(interp.get("max_daily_range")): text += f"；最大日变化幅度为{_value_unit(interp['max_daily_range'], unit)}，出现在{pd.Timestamp(interp['max_daily_range_date']).strftime('%Y年%m月%d日')}"
        paragraphs.append(text + "。")
    return paragraphs


def generate_single_variable_report(context):
    """Generate a non-empty .docx report from a validated report context."""
    if not context.get("software_info", {}).get("qc_confirmed"):
        raise ValueError("请先确认最终质控结果，再生成报告。")
    doc = Document(); section = doc.sections[0]; section.page_width = Cm(21); section.page_height = Cm(29.7); section.left_margin = section.right_margin = Cm(2.2)
    styles = doc.styles; styles["Normal"].font.name = "宋体"; styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体"); styles["Normal"].font.size = Pt(10.5); styles["Normal"].paragraph_format.first_line_indent = Pt(0); styles["Normal"].paragraph_format.left_indent = Pt(0); styles["Normal"].paragraph_format.right_indent = Pt(0)
    footer = section.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.CENTER; _page_number(footer)
    project, variable, data, qc, stats, features, params, software = (context[k] for k in ("project_info", "variable_info", "data_summary", "qc_summary", "statistics", "time_features", "parameters", "software_info"))
    name, unit = variable["display_name"], variable["unit"]

    doc.add_paragraph(project["report_title"], style="Title").alignment = WD_ALIGN_PARAGRAPH.CENTER
    for label, value in (("站点名称", project.get("site_name", "")), ("项目名称", project.get("project_name", "")), ("当前变量", f"{name}（{unit}）"), ("监测时间范围", f"{_date(data['start_time'])} 至 {_date(data['end_time'])}"), ("编制部门", project.get("department", "")), ("编制人", project["author"]), ("报告生成日期", _date(software["generated_at"])), ("业务内核版本", software["software_version"]), ("业务基线标签", software["git_tag"])):
        if value: doc.add_paragraph(f"{label}：{value}")
    doc.add_page_break()

    _add_heading(doc, "1. 数据摘要")
    summary = f"本次{name}监测数据共包含{data['raw_count']}条记录，监测时间为{_date(data['start_time'])}至{_date(data['end_time'])}，日历覆盖{_fmt(data.get('calendar_days'), 0)}天，实际有数据{_fmt(data.get('actual_data_days'), 0)}天，中位采样间隔为{_fmt(data['median_sampling_interval_minutes'], 2)}分钟。经自动质控和人工复核后，最终保留{data['final_valid_count']}条有效记录，最终有效率为{_fmt(data['valid_rate'], 2)}%。"
    _format_body_paragraph(doc.add_paragraph(summary), 10.5)
    _table(doc, "表1 数据摘要", ["项目", "结果"], [("变量名称", name), ("单位", unit), ("监测开始时间", _date(data["start_time"])), ("监测结束时间", _date(data["end_time"])), ("日历覆盖天数", _fmt(data.get("calendar_days"), 0)), ("实际有数据天数", _fmt(data.get("actual_data_days"), 0)), ("日期覆盖率", f"{_fmt(data.get('date_coverage_rate'), 2)}%"), ("中位采样间隔（分钟）", _fmt(data["median_sampling_interval_minutes"], 2)), ("原始记录数", data["raw_count"]), ("最终有效记录数", data["final_valid_count"]), ("最终有效率", f"{_fmt(data['valid_rate'], 2)}%")])

    _add_heading(doc, "2. 数据质控与前后对比"); _format_body_paragraph(doc.add_paragraph("本次数据依次进行了原始缺测识别、传感器零值处理、硬范围检查、Hampel异常候选识别、连续恒定值识别和人工复核。传感器零值及超出硬范围的记录自动剔除，Hampel和恒定值结果仅作为候选，由人工确认后形成最终质控数据。"), 10.5)
    _format_body_paragraph(doc.add_paragraph(_qc_text(data, qc)), 10.5)
    _table(doc, "表2 质控结果", ["质控项目", "数量", "占原始记录比例", "处理方式"], [("原始缺测", data["raw_missing_count"], _pct(data["raw_missing_count"], data["raw_count"]), "识别并保留为缺测"), ("传感器0值", qc.get("removed_by_sensor_zero", 0), _pct(qc.get("removed_by_sensor_zero", 0), data["raw_count"]), "自动剔除"), ("超出硬范围", qc.get("removed_by_range", 0), _pct(qc.get("removed_by_range", 0), data["raw_count"]), "自动剔除"), ("Hampel候选", qc.get("flagged_by_hampel", 0), _pct(qc.get("flagged_by_hampel", 0), data["raw_count"]), "候选标记，人工复核"), ("恒定值候选", qc.get("flagged_by_constant_value", 0), _pct(qc.get("flagged_by_constant_value", 0), data["raw_count"]), "候选标记，人工复核"), ("人工删除", qc.get("manual_remove_count", 0), _pct(qc.get("manual_remove_count", 0), data["raw_count"]), "人工复核决策"), ("最终缺测", data["final_missing_count"], _pct(data["final_missing_count"], data["raw_count"]), "最终结果"), ("最终有效记录", data["final_valid_count"], _pct(data["final_valid_count"], data["raw_count"]), "用于后续分析")])

    source = context["source_data"]; _add_picture(doc, _comparison_figure(source["raw_data"], source["final_qc_data"], name, unit), f"图1 {name}原始数据与最终质控数据对比（上下图纵轴范围分别自适应）"); _format_body_paragraph(doc.add_paragraph(_comparison_text(source, features.get("interpretation", {}), unit)), 10.5)

    _add_heading(doc, "3. 统计特征与时间变化"); _format_body_paragraph(doc.add_paragraph("基础统计基于最终质控后的小时平均数据计算。"), 10.5)
    interp = features.get("interpretation", {})
    _table(doc, "表3 统计特征", ["指标", "数值", "单位"], [("最小值", _fmt(stats.get("min")), unit), ("最大值", _fmt(stats.get("max")), unit), ("平均值", _fmt(stats.get("mean")), unit), ("中位数", _fmt(stats.get("median")), unit), ("标准差", _fmt(stats.get("std")), unit), ("P5", _fmt(interp.get("p5")), unit), ("P95", _fmt(interp.get("p95")), unit), ("主体90%范围", f"{_fmt(interp.get('p5'))} 至 {_fmt(interp.get('p95'))}", unit), ("日变化幅度中位数", _fmt(interp.get("daily_range_median")), unit), ("日变化幅度P95", _fmt(interp.get("daily_range_p95")), unit), ("最大日变化幅度", _fmt(interp.get("max_daily_range", stats.get("max_daily_range"))), unit)])

    doc.add_paragraph("时间变化综合分析", style="Heading 2")
    narratives = context.get("narratives", {})
    for narrative_key in ("hourly", "intraday", "monthly"):
        if narratives.get(narrative_key):
            _format_body_paragraph(doc.add_paragraph(narratives[narrative_key]), 10.5)

    interpretation = features.get("interpretation", {}); mode_text = {"trend_up": "总体呈上升趋势", "trend_down": "总体呈下降趋势", "decrease_then_increase": "呈先下降后上升变化", "increase_then_decrease": "呈先上升后下降变化", "stable": "总体较为稳定", "change_point": f"在{_month(interpretation.get('change_time')) if interpretation.get('change_time') else '观测期内'}前后出现明显水平{MODE_LABELS.get(interpretation.get('change_direction'), '变化')}"}.get(interpretation.get("narrative_mode"), "总体较为稳定")
    # The three context narratives above are the sole report analysis block and precede all figures.
    # Hourly extremes and the P5–P95 range are rendered once after the time-series figures.
    resampled, anomaly, monthly = features.get("resampled", {}), features.get("anomaly"), stats.get("monthly")
    hourly, daily = resampled.get("hourly", pd.DataFrame()), resampled.get("daily", pd.DataFrame())
    if not hourly.empty and not daily.empty: _add_picture(doc, _hourly_daily_figure(hourly, daily, name, unit), f"图2 {name}小时平均与日平均")
    if anomaly is not None and not anomaly.empty: _add_picture(doc, _anomaly_figure(anomaly, name, unit), f"图3 {name}日内距平")
    if monthly is not None and not monthly.empty: _add_picture(doc, _monthly_figure(monthly, name, unit), f"图4 {name}月平均与月标准差")
    # Do not add analysis text between or after the three figures.
    if monthly is not None and not monthly.empty:
        valid_monthly = monthly.dropna(subset=["monthly_mean"])
        if not valid_monthly.empty:
            high, low = valid_monthly.loc[valid_monthly.monthly_mean.idxmax()], valid_monthly.loc[valid_monthly.monthly_mean.idxmin()]
            pass
        valid_std = monthly.dropna(subset=["monthly_std"])
        if not valid_std.empty:
            pass
    ranges = stats.get("daily_range")
    if ranges is not None and not ranges.empty and ranges["daily_range"].notna().any():
        pass

    _add_heading(doc, "4. 结论"); _format_body_paragraph(doc.add_paragraph(f"本次{name}监测数据原始记录数为{data['raw_count']}条，已完成自动质控与人工复核，最终保留{data['final_valid_count']}条有效记录，有效率为{_fmt(data['valid_rate'], 2)}%。传感器零值处理{qc.get('removed_by_sensor_zero', 0)}条，超出硬范围处理{qc.get('removed_by_range', 0)}条，人工删除{qc.get('manual_remove_count', 0)}条，相关过程可通过QC日志追溯。"), 10.5)
    main_fact = {"change_point": f"{_month(interpretation['change_time']) if interpretation.get('change_time') else '观测期内'}前后出现明显水平{MODE_LABELS.get(interpretation.get('change_direction'), '变化')}", "trend_up": "总体呈上升趋势", "trend_down": "总体呈下降趋势", "decrease_then_increase": "呈先下降后上升变化", "increase_then_decrease": "呈先上升后下降变化"}.get(interpretation.get("narrative_mode"), "总体较为稳定")
    primary = interpretation.get("primary_mode", interpretation.get("narrative_mode"))
    monthly_info = interpretation.get("monthly_interpretation", {})
    if narratives.get("conclusion"):
        conclusion = narratives["conclusion"]
    elif primary == "periodic":
        conclusion = f"观测期内日内距平呈规律性重复波动，日变化幅度中位数为{_value_unit(interpretation.get('daily_range_median'), unit)}。"
        if monthly_info.get("monthly_min_month") is not None: conclusion += f"月平均变化方面，{_month(monthly_info['monthly_min_month'])}为观测期低值阶段。"
    elif primary == "change_point":
        conclusion = f"观测期内{main_fact}。"
        if monthly_info.get("pre_mean") is not None: conclusion += f"转折前后月平均水平由{_value_unit(monthly_info['pre_mean'], unit)}变为{_value_unit(monthly_info['post_mean'], unit)}。"
        if monthly_info.get("monthly_min_month") is not None: conclusion += f"{_month(monthly_info['monthly_min_month'])}月平均值为观测期最低。"
    else:
        conclusion = f"观测期内该变量{main_fact}。"
        if monthly_info.get("monthly_max_month") is not None: conclusion += f"月平均最高值出现在{_month(monthly_info['monthly_max_month'])}。"
    _format_body_paragraph(doc.add_paragraph(conclusion), 10.5)

    _add_heading(doc, "附录A：质控参数"); _table(doc, "表A-1 质控参数", ["参数", "值"], [("sensor_zero 是否启用", "是" if params["sensor_zero_enabled"] else "否"), ("hard_min", _fmt(params["hard_min"])), ("hard_max", _fmt(params["hard_max"])), ("Hampel窗口", _fmt(params["hampel_window"])), ("Hampel阈值", _fmt(params["hampel_threshold"])), ("Hampel最小绝对偏差", _fmt(params["hampel_min_abs_deviation"])), ("恒定值窗口", _fmt(params["constant_value_window"])), ("恒定值容差", _fmt(params["constant_value_tolerance"]))])
    _add_heading(doc, "附录B：处理信息"); _table(doc, "表B-1 处理信息", ["项目", "值"], [("variable_key", variable["variable_key"]), ("数据时间范围", f"{_date(data['start_time'])} 至 {_date(data['end_time'])}"), ("当前人工确认状态", "已确认"), ("报告生成时间", _date(software["generated_at"])), ("业务内核版本", software["software_version"]), ("业务基线标签", software["git_tag"])])
    _add_heading(doc, "附录C：日志说明"); _format_body_paragraph(doc.add_paragraph("完整逐点质控记录请参见系统导出的 final_qc_log.xlsx。"), 10.5)
    _format_body_paragraph(doc.add_paragraph("说明：传感器精确0值按照当前项目规则作为无效值处理；hard_range 按项目变量配置执行。"), 10.5)
    output = BytesIO(); doc.save(output); return output.getvalue()
